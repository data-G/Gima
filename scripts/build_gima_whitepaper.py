#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "GIMA_WHITEPAPER.md"
OUTPUT = ROOT / "outputs" / "whitepaper"
DOCX = OUTPUT / "Gima_Local_First_AI_White_Paper.docx"
DIAGRAM = OUTPUT / "gima_architecture.png"

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0F766E"
PALE = "E8EEF5"
LIGHT = "F2F4F7"
MUTED = "5D6875"
WHITE = "FFFFFF"
INK = "20262E"
GOLD = "B07A1B"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def create_architecture_diagram(path: Path) -> None:
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 38)
        label_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 25)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 21)
    except OSError:
        title_font = label_font = body_font = ImageFont.load_default()
    draw.text((55, 35), "Gima local-first architecture", font=title_font, fill="#17324D")
    boxes = [
        (55, 120, 330, 285, "Interfaces", "Web UI\nCLI and voice\nFiles and prompts", "#E8EEF5"),
        (385, 120, 690, 285, "Routing", "Brain-first search\nDeterministic tasks\nLocal/teacher models", "#DDF3F0"),
        (745, 120, 1040, 285, "Production", "Documents and sheets\nCode and tools\nAudio and video", "#FFF2D9"),
        (1095, 120, 1445, 285, "User outputs", "Download cards\nHands/out projects\nReviewable manifests", "#E8EEF5"),
        (215, 420, 570, 640, "Local memory", "CSV knowledge and conversations\nRebuildable FTS index\nBrain files and source reviews", "#EEF4FA"),
        (625, 420, 980, 640, "Governance", "Scoped permissions and secrets\nAllowlisted/sandboxed tools\nApproval-gated self-update", "#F9E8E8"),
        (1035, 420, 1390, 640, "Continuity", "Daily review learning\nSource and state snapshots\nTests, logs, and rollback", "#E9F5E9"),
    ]
    for x1, y1, x2, y2, heading, body, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=fill, outline="#8FA3B8", width=3)
        draw.text((x1 + 24, y1 + 20), heading, font=label_font, fill="#17324D")
        draw.multiline_text((x1 + 24, y1 + 70), body, font=body_font, fill="#273746", spacing=10)
    for start, end in [((330, 202), (385, 202)), ((690, 202), (745, 202)), ((1040, 202), (1095, 202))]:
        draw.line((*start, *end), fill="#0F766E", width=7)
        draw.polygon([(end[0], end[1]), (end[0]-18, end[1]-11), (end[0]-18, end[1]+11)], fill="#0F766E")
    for x in (510, 835, 1215):
        draw.line((x, 285, x, 405), fill="#8295A8", width=5)
        draw.polygon([(x, 420), (x-11, 401), (x+11, 401)], fill="#8295A8")
    draw.text((55, 700), "Online AI providers are optional teachers; local storage remains the durable system of record.", font=body_font, fill="#5D6875")
    image.save(path, quality=95)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "GIMA WHITE PAPER"
    set_font(header.runs[0], size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(82)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_font(kicker.add_run("TECHNICAL WHITE PAPER"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("Gima"), size=34, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("A Local-First, Review-Gated\nPersonal AI Workspace"), size=20, color=DARK_BLUE, bold=True)
    deck = doc.add_paragraph()
    deck.alignment = WD_ALIGN_PARAGRAPH.CENTER
    deck.paragraph_format.space_after = Pt(74)
    set_font(deck.add_run("Architecture, continuous learning, multimodal production,\nand human-controlled self-improvement"), size=12, color=MUTED, italic=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    set_font(meta.add_run("Version 1.0 | 20 June 2026"), size=11, color=NAVY, bold=True)
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(author.add_run("Gimhan Gunarathne | Gima Project"), size=10.5, color=MUTED)
    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.paragraph_format.space_before = Pt(40)
    set_font(repo.add_run("github.com/data-G/Gima"), size=10, color=TEAL, bold=True)
    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for heading in headings:
        para = doc.add_paragraph(style="List Bullet")
        set_font(para.add_run(re.sub(r"^\d+\.\s*", "", heading)), size=10.5, color=INK)
    doc.add_page_break()


def add_inline_markup(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    if cols == 3:
        widths = [1900, 3300, 4160]
    elif cols == 2:
        widths = [2700, 6660]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markup(paragraph, value.strip())
            for run in paragraph.runs:
                run.font.size = Pt(9.5)
            if r_index == 0:
                set_cell_shading(cell, LIGHT)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> Path:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    create_architecture_diagram(DIAGRAM)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    headings = [line[3:].strip() for line in lines if line.startswith("## ") and line[3:4].isdigit()]
    add_contents(doc, headings)

    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        if not line or line.startswith("# Gima:") or line.startswith("**Architecture") or line.startswith("Version 1.0") or line.startswith("Gimhan ") or line.startswith("Repository:"):
            index += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if first_heading:
                first_heading = False
            doc.add_paragraph(heading, style="Heading 1")
            if heading == "3. System Architecture":
                doc.add_picture(str(DIAGRAM), width=Inches(5.7))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                set_font(caption.add_run("Figure 1. Gima's layered local-first architecture."), size=9, color=MUTED, italic=True)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("| "):
            table_rows = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", cell) for cell in cells):
                    table_rows.append(cells)
                index += 1
            add_markdown_table(doc, table_rows)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markup(paragraph, line[2:].strip())
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markup(paragraph, re.sub(r"^\d+\. ", "", line))
            index += 1
            continue
        paragraph = doc.add_paragraph()
        add_inline_markup(paragraph, line)
        index += 1

    doc.core_properties.title = "Gima: A Local-First, Review-Gated Personal AI Workspace"
    doc.core_properties.subject = "Technical white paper"
    doc.core_properties.author = "Gimhan Gunarathne"
    doc.core_properties.keywords = "Gima, local-first AI, personal AI, continuous learning, responsible AI"
    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    print(build())
